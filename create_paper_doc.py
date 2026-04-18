from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('An Age–Period–Cohort Framework for Profit and Profit Volatility Modeling', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author / Source
p = doc.add_paragraph('Joseph L. Breeden')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Published in Mathematics (MDPI), Vol. 12, No. 10, Article 1427, 2024')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

sections = [
    ("Abstract / Overview",
     "This paper presents a Framework for Profit and Uncertainty (FPU) modeling that unifies credit scoring, "
     "finance, and predictive analytics through discrete-time survival modeling and age-period-cohort (APC) "
     "approaches. The framework addresses critical lending challenges including machine learning integration, "
     "forecasting uncertainty quantification, climate risk stress testing, portfolio optimization, and normalizing "
     "pandemic-era data for credit modeling."),

    ("1. Introduction",
     "The lending industry has experienced significant technological advancement through machine learning and "
     "alternative data sources, yet these innovations often fail to address fundamental challenges. The largest "
     "portfolio failures come from systematic shifts in the probability of default relative to a credit score or "
     "an inability to price loans accurately.\n\n"
     "Traditional credit scoring architectures present a critical structural mismatch. Credit scores utilize "
     "cross-sectional data where accounts appear once with a binary outcome indicator across fixed observation "
     "windows of 12–48 months. This approach cannot capture timing dynamics or adjust for macroeconomic "
     "scenarios essential to yield prediction.\n\n"
     "Survival modeling provides an integrated framework. The paper presents a Framework for Profit and "
     "Uncertainty (FPU) modeling that synthesizes existing research into an integrated framework applicable to "
     "finance, credit risk, and loan origination departments."),

    ("2. Start from the Goal and Work Backward",
     "Legacy frameworks in lending derive from 1960s paper scorecard methodologies. The fundamental business "
     "objective differs substantially from current practice: the goal of lending is not to rank-order risk. "
     "Lenders provide money to borrowers with the expectation of achieving a positive net yield above their "
     "established hurdle rate.\n\n"
     "Contemporary lending faces compressed margins. Declining yields from price competition and regulatory "
     "compliance costs leave minimal room for error in product design and pricing. Effective lending requires "
     "forward-looking cash flow models incorporating account lifecycle functions (hazard rates), macroeconomic "
     "scenarios, borrower attribute risk adjustments, and segment-level risk modifications.\n\n"
     "Age-period-cohort models provide natural alignment with these requirements while integrating account-level scores."),

    ("3. Age–Period–Cohort Models as a Credit Risk Framework",
     "Age-period-cohort models and survival models share mathematical foundations. Cox proportional hazards "
     "models extend survival analysis to estimate relative risk between individuals and cohorts, with "
     "applications in credit scoring.\n\n"
     "Core Challenge: Cox proportional hazards models accommodate only two dimensions effectively. Credit risk "
     "modeling requires three critical dimensions: age of account (a), origination date/vintage (v), and "
     "calendar date (t). The mathematical relationship a = t − v creates linear dependency, potentially "
     "producing unstable and biased coefficient estimates.\n\n"
     "APC Model Formulation — For probability of default modeling with binomial distribution:\n"
     "    logit(PD) = F(a) + G(v) + H(t)\n"
     "Where F(a) = lifecycle function, G(v) = vintage function, H(t) = environment function.\n\n"
     "Decomposition with Linear and Nonlinear Components:\n"
     "    logit(PD) = α₀ + α₁a + F′(a) + β₁v + G′(v) + γ₁t + H′(t)\n\n"
     "For data spanning multiple economic cycles, standard practice assumes γ₁ = 0, consistent with "
     "through-the-cycle (TTC) average probability of default central to Basel II regulatory capital requirements. "
     "Empirical testing confirms APC formulations demonstrate 1–3% in-sample errors over one-year periods for "
     "sufficiently large datasets."),

    ("4. Profit Models",
     "Cash flow modeling aggregates forecast components and loan terms to calculate expected yield. Key components "
     "include the probability of active account P_Active(a), revenue (scheduled payments, fees, interest), costs "
     "(servicing, cost of funds), and loss recovery upon default.\n\n"
     "Each month, accounts face three possible states: default (PD), prepayment (PA), or continuation. Both PD "
     "and PA are conditional on prior-month activity status.\n\n"
     "    Total Profit = Σ[Incremental Cash Flow(t)] + [Delinquency/Charge-off/Recovery Periods]\n\n"
     "Account-level pricing is impossible without account-level yield forecasting. APC frameworks provide monthly "
     "forecasts of conditional PD and PA by segment and vintage. Replacing the vintage function G(v) with "
     "attribute-based scores enables account-level yield forecasts."),

    ("5. Credit Scoring for Cash Flow Forecasting",
     "Rank-order credit scores misalign with forward-looking probability of default estimation required for cash "
     "flow modeling. Panel data structures enable survival model approaches where timing connects to exogenous "
     "events and account maturation.\n\n"
     "Panel Data Credit Score with APC Framework:\n"
     "    logit(PD(i,a,t)) = F(a) + H(t) + C(i)\n"
     "Where F(a) + H(t) are fixed offsets (previously estimated APC components) and C(i) = Σⱼ cⱼ sᵢⱼ is "
     "the credit score component.\n\n"
     "The scoring component C(i) remains in log-odds units, paralleling traditional credit scores while enabling "
     "integration into cash flow models. APC credit scores demonstrate greater out-of-sample robustness while "
     "maintaining coefficients generally similar to traditional cross-sectional scores."),

    ("6. Loss Reserves and Economic Capital",
     "Credit scores from panel data function across any forecast horizon. The scoring component C(i) remains "
     "independent of horizon, enabling single-model application to multiple time intervals.\n\n"
     "IFRS 9 Loss Reserve Application: IFRS 9 requires stage-based loss reserve methodology:\n"
     "  • Stage 1: 12-month loss reserves (performing accounts)\n"
     "  • Stage 2: Lifetime loss reserves (deteriorated accounts, ~20% of portfolio)\n"
     "  • Stage 3: Default with post-default accounting treatment\n\n"
     "The APC solution uses a single model with different summation intervals — Stage 1 sums 12-month forecasts; "
     "Stage 2 sums lifetime forecasts. Reserves adjust smoothly as accounts age through the lifecycle function F(a).\n\n"
     "Economic Capital: Capital allocation proportional to relative economic capital enables scaling cost of "
     "capital by segment risk needs. Loss volatility (unexpected loss divided by expected loss) varies "
     "significantly by credit quality: subprime borrowers have a UL/EL ratio of approximately 0.5 while "
     "superprime borrowers have a ratio of approximately 4.\n\n"
     "Prepayment Scores: Prepayment modeling achieves parity with default modeling in importance, particularly "
     "for prime lending where voluntary prepayment may be the greatest risk of unprofitable lending.\n\n"
     "    Unconditional Default = Conditional PD × P_Active(previous period)\n"
     "    Unconditional Prepayment = Conditional PA × P_Active(previous period)"),

    ("7. Loss Recovery Prediction",
     "Recovery from charged-off debt contributes materially to cash flow analysis. Recovery processes may extend "
     "years, preventing simple moving-average calculations for young portfolios.\n\n"
     "APC models apply effectively using charge-off date as the initial event. The lifecycle function F(a) "
     "measures recovery probability versus defaulted debt age. Expected incremental recoveries condition on "
     "original charge-off probability, are discounted for net present value, and are added to total cash flow."),

    ("8. The Role of Machine Learning and AI",
     "Machine learning adoption in credit risk has expanded dramatically. Gradient-boosted regression trees and "
     "ReLU neural networks demonstrate strongest empirical performance and interpretability.\n\n"
     "Fundamental Limitation: 'The root cause of the limitations of ML, AI, and generative AI in credit risk is "
     "the sparsity of data. We may have many accounts to model, but not enough economic cycles or credit cycles "
     "for any statistical algorithm to predict borrower behavior without a conceptual understanding of the "
     "borrowers, the loan products, and the economy.'\n\n"
     "APC + Machine Learning Integration: Rather than replacing survival models with machine learning, combining "
     "both approaches leverages complementary strengths. Stochastic Gradient Boosted Regression Tree (SGBRT) "
     "libraries in R and Python support 'offset' specification — the offset includes previously-estimated APC "
     "components without scaling coefficients, leaving machine learning to explain residuals.\n\n"
     "A dual-network neural architecture is used where APC inputs (lifecycle + environment) flow to output "
     "unmodified and a ReLU network processes account attributes; both combine for the total PD forecast. "
     "Integrated APC + ML models demonstrate substantially greater out-of-sample robustness."),

    ("9. Incorporating Alternate Data",
     "Machine learning demonstrates greatest advantages in niche products, alternate channels, and underbanked "
     "populations utilizing non-traditional data sources: deposit histories, corporate financial statements, "
     "social media, and mobile phone usage patterns.\n\n"
     "The APC framework separates account idiosyncratic effects from systemic environment effects, improving "
     "signal-to-noise ratios in alternate data, mitigating spurious correlations, and preventing double-counting "
     "against economic environment factors."),

    ("10. Understanding Credit Cycles",
     "Credit cycles occur when loans from specific periods demonstrate better or worse intrinsic risk even after "
     "normalizing for scoring factors and post-origination economic conditions.\n\n"
     "Two distinct adverse selection mechanisms exist:\n"
     "  • Competitive Adverse Selection: Price-based lender differentiation attracts higher-risk borrowers. "
     "Generally non-cyclical.\n"
     "  • Macroeconomic Adverse Selection: Autocorrelated over years; a 'hidden variable' problem arising from "
     "borrower personality differences not captured in credit bureau or application data.\n\n"
     "Measurement: APC origination scores isolate vintage residuals revealing macroeconomic adverse selection. "
     "These residuals correlate with Senior Loan Officer Opinion Survey data on mortgage demand. Analysis of "
     "Fannie Mae/Freddie Mac data shows the highest score bands have the greatest cycle dynamics, with "
     "2018–2019 quality deterioration concentrated in high-score populations.\n\n"
     "Early Warning: Adverse selection measurement through panel data demonstrates early warning capability "
     "before performance materialization, with traditional cross-sectional methods showing a 6–9 month lag."),

    ("11. Incorporating Economic Cycles",
     "Forward-looking economic scenarios are essential for competitive loan origination and portfolio management. "
     "APC and survival model stress-testing applications have long been deployed for scenario analysis.\n\n"
     "Direct scenario integration restructures pricing methodology — management's economic stance flows "
     "immediately through portfolio and account-level profitability estimates. Economic momentum enables "
     "reasonable predictability approximately 12 months forward.\n\n"
     "For loss forecasting, 24-month scenario windows capture near-term recession/expansion impacts. Beyond "
     "24 months, reversion to long-run averages reflects genuine uncertainty and aligns with 'reasonable and "
     "supportable' periods used by most lenders."),

    ("12. Profit Volatility Models",
     "Portfolio optimization requires understanding not only expected profitability but also profit volatility. "
     "Modern Portfolio Theory optimization applies to lending through the Sharpe Ratio:\n\n"
     "    Sharpe Ratio = (Expected Return − Risk-Free Rate) / Volatility of Return\n\n"
     "Yield uncertainty derives from three sources: lifecycle function estimation uncertainty, credit risk "
     "estimation uncertainty, and future environment forecast uncertainty.\n\n"
     "Monte Carlo experiments consistently produce beta distribution fits with two parameters; fifty iterations "
     "sufficiently estimate lifetime yield uncertainty, enabling one-second laptop computation suitable for "
     "portfolio optimization.\n\n"
     "APC data decomposition extracting environment functions followed by correlation analysis among these "
     "functions provides appropriate correlation matrices in log-odds space suitable for normal distribution assumptions."),

    ("13. Optimizing Pricing for Net Income",
     "'Meet the market' pricing frequently fails through creating herd mentality competing for growth without "
     "profitability prediction. Yield forecasts enable restructured pricing methodology.\n\n"
     "    Net Income = (Expected Yield × Volume) − (Operating Costs × Volume)\n\n"
     "Interest rate increases enhance estimated yield while reducing volume; rate decreases lower yield but "
     "increase volume. Natural optimal loan term values emerge from this tradeoff. Single portfolio policies "
     "rarely outperform market equilibrium — segment-level optimization identifies missed opportunities, "
     "particularly during economic transitions."),

    ("14. Underwriting with Profitability and Profit Uncertainty Estimates",
     "Portfolio failures rarely result from rank-ordering inability; failures occur through inadequate score "
     "cutoff management across changing economic conditions ensuring profitable lending.\n\n"
     "APC and discrete-time survival models enable credit risk component separation from lifecycle and "
     "environment. Score cutoffs derive immediately from lifetime loss or profitability goals. Cash flow model "
     "targets adjust thresholds to achieve goals under chosen economic scenarios.\n\n"
     "Uncertainty-Informed Underwriting: Forecast uncertainty varies substantially across input space. Low scores "
     "from thin-file borrowers indicate genuine uncertainty; low scores from thick-file borrowers indicate "
     "confidence in poor credit risk. Panel data APC models enable efficient full forecast uncertainty estimation "
     "through Monte Carlo simulation estimating beta distribution parameters per borrower — despite computational "
     "simplicity, this uncertainty information remains unused in lending, a significant missed opportunity."),

    ("15–17. Stress Testing, Pandemic Normalization, and Framework Integration",
     "The framework continues addressing:\n"
     "  • Stress Testing: Extreme event incorporation including climate change and pandemic scenarios.\n"
     "  • Pandemic Era Data Normalization: APC adjustments enabling pandemic-period data in scoring models.\n"
     "  • Framework Integration: A unified approach consolidating all components into a single coherent system."),

    ("Conclusion",
     "This integrated Framework for Profit and Uncertainty (FPU) unifies previously separate lending functions — "
     "origination, pricing, portfolio management, risk measurement — through common mathematical foundations.\n\n"
     "By combining age-period-cohort decomposition with contemporary machine learning and incorporating full "
     "probability forecasting rather than rank-ordering, lenders can achieve more profitable, risk-appropriate, "
     "and economically-responsive lending practices.\n\n"
     "Key advantages include: single models serving multiple functions across time horizons; uncertainty "
     "quantification informing underwriting; natural integration of stress testing; and tractable portfolio "
     "optimization with genuine profit-risk tradeoffs."),
]

for heading, body in sections:
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.add_paragraph()

# Source note
doc.add_paragraph('─' * 60)
note = doc.add_paragraph(
    'Source: Breeden, J.L. (2024). An Age–Period–Cohort Framework for Profit and Profit Volatility Modeling. '
    'Mathematics, 12(10), 1427. https://doi.org/10.3390/math12101427\n'
    'Full paper: https://www.mdpi.com/2227-7390/12/10/1427'
)
note.runs[0].italic = True

output_path = r'C:\Users\mchad\OneDrive\Documents\Mitch\APC_Profit_Volatility_Paper.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
